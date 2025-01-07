import docx
import os
from pathlib import Path
from django.core.files.storage import FileSystemStorage

# Build paths inside the project like this: BASE_DIR / 'subdir'.
BASE_DIR = Path(__file__).resolve().parent.parent

# Create your models here.
def get_text(filename, search):
    doc = docx.Document(filename)
    fullText = []
    for para in doc.paragraphs:
        if search.strip().lower() in para.text.lower():
            fullText.append(para.text)
    return '\n'.join(fullText)


def full_text_data(filename):
    doc = docx.Document(filename)
    fullText = ""
    for para in doc.paragraphs:
        fullText += para.text
    return fullText

def full_text(filename):
    doc = docx.Document(filename)
    fullText = ""
    for para in doc.paragraphs:
        if para.text != '':
            fullText += para.text
    return fullText

def check_files(file_name):
    try:
        files = os.listdir(os.path.join(BASE_DIR, 'media'))
        files_path = os.path.join(BASE_DIR, 'media')
        is_exist = False
        if file_name in files:
            is_exist = True
        return is_exist
    except Exception as error:
        print(error)


def full_details_pending():
    try:
        files = os.listdir(os.path.join(BASE_DIR, 'media'))
        files_path = os.path.join(BASE_DIR, 'media')
        arr = []
        for file in files:
            if 'pending' in file:
                docx_path = files_path + '\\' + file
                fs = FileSystemStorage()
                file_path = fs.url(file)
                text = full_text(docx_path)
                if text != "":
                    file_dict = {}
                    file_dict["file"] = file
                    file_dict["path"] = file_path
                    file_dict["content"] = text
                    arr.append(file_dict)
        return arr
    except Exception as error:
        print(error)


def full_details_approved():
    try:
        files = os.listdir(os.path.join(BASE_DIR, 'media'))
        files_path = os.path.join(BASE_DIR, 'media')
        arr = []
        i = 0;
        for file in files:
            if 'approved' in file:
                docx_path = files_path + '\\' + file
                fs = FileSystemStorage()
                file_path = fs.url(file)
                text = full_text(docx_path)
                text_data = full_text_data(docx_path)
                if text != "":
                    file_dict = {}
                    file_dict["file"] = file
                    file_dict["path"] = file_path
                    file_dict["content"] = text
                    file_dict["text"] = text_data
                    file_dict["number"] = i
                    arr.append(file_dict)
        return arr
    except Exception as error:
        print(error)


def get_research(search):
    try:
        files = os.listdir(os.path.join(BASE_DIR, 'media'))
        files_path = os.path.join(BASE_DIR, 'media')
        arr = []
        i = 0;
        for file in files:
            
            if 'approved' in file:
                docx_path = files_path + '\\' + file
                fs = FileSystemStorage()
                file_path = fs.url(file)
                text = get_text(docx_path, search)
                print(text)
                text_data = full_text_data(docx_path)
                
                if text != "":
                    i = i + 1
                    file_dict = {}
                    file_dict["file"] = file.replace("approved", " ")
                    file_dict["path"] = file_path
                    file_dict["content"] = text
                    file_dict["text"] = text_data
                    file_dict["number"] = i
                    arr.append(file_dict)
        return arr
    except Exception as error:
        print(error) 
        
# def get_research(search):
#     try:
#         files = os.listdir("D:/abstract-repository/abstractrepoistory/abstract")
#         print(files)
#         # os.chdir('abstract')
#         arr = []
#         for file in files:
#             docx_path = "D:/abstract-repository/abstractrepoistory/abstract/" + file
#             if search in getText(docx_path):
#                 arr.append(getText(docx_path))
#         return arr
#     except Exception as error:
#         print(error) 
    
    